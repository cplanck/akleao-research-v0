"""Document ingestion module - handles loading various document types."""

from pathlib import Path
from dataclasses import dataclass
from typing import Iterator
import hashlib
import tempfile
import requests
import os
from urllib.parse import urlparse


@dataclass
class TextSegment:
    """A segment of text with page information."""
    text: str
    page_number: int | None = None
    section: str | None = None


@dataclass
class Document:
    """Represents a loaded document."""
    content: str
    source: str
    doc_type: str
    metadata: dict = None
    # For page-aware documents, this contains segments with page numbers
    segments: list[TextSegment] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
        if self.segments is None:
            self.segments = []
        # Generate a unique ID based on content hash
        self.id = hashlib.md5(self.content.encode()).hexdigest()[:12]


class DocumentLoader:
    """Loads documents from various sources."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".markdown"}

    def __init__(self, fast_mode: bool = False):
        """Initialize document loader.

        Args:
            fast_mode: If True, skip table structure recognition for faster PDF processing.
                       Good for text-heavy PDFs without complex tables.
        """
        self.fast_mode = fast_mode

    def load(self, path: str | Path) -> list[Document]:
        """Load document(s) from a file or directory path."""
        path = Path(path)

        if path.is_dir():
            return list(self._load_directory(path))
        elif path.is_file():
            return [self._load_file(path)]
        else:
            raise FileNotFoundError(f"Path not found: {path}")

    def _load_directory(self, dir_path: Path) -> Iterator[Document]:
        """Recursively load all supported documents from a directory."""
        for file_path in dir_path.rglob("*"):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    yield self._load_file(file_path)
                except Exception as e:
                    print(f"Warning: Failed to load {file_path}: {e}")

    def _load_file(self, file_path: Path) -> Document:
        """Load a single document file."""
        suffix = file_path.suffix.lower()
        segments = []

        if suffix == ".pdf":
            content, segments = self._load_pdf(file_path)
            doc_type = "pdf"
        elif suffix == ".docx":
            content = self._load_docx(file_path)
            doc_type = "docx"
        elif suffix in {".md", ".markdown"}:
            content = self._load_text(file_path)
            doc_type = "markdown"
        elif suffix == ".txt":
            content = self._load_text(file_path)
            doc_type = "text"
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        return Document(
            content=content,
            source=str(file_path),
            doc_type=doc_type,
            metadata={"filename": file_path.name},
            segments=segments
        )

    def _load_pdf(self, file_path: Path) -> tuple[str, list[TextSegment]]:
        """Extract text from PDF using Docling for better table/image handling.

        Returns:
            tuple of (content, segments_with_page_numbers)
        """
        try:
            return self._load_pdf_with_docling(file_path)
        except Exception as e:
            print(f"Docling failed, falling back to pypdf: {e}")
            # Basic fallback doesn't have page-level segments
            content = self._load_pdf_basic(file_path)
            return content, []

    def _load_pdf_basic(self, file_path: Path) -> str:
        """Basic PDF extraction using pypdf (fallback)."""
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts)

    def _load_pdf_with_docling(self, file_path: Path) -> tuple[str, list[TextSegment]]:
        """Extract text from PDF using Docling - handles tables and images.

        Returns:
            tuple of (full_markdown_content, list_of_segments_with_page_numbers)
        """
        from docling.document_converter import DocumentConverter
        from docling.datamodel.base_models import InputFormat
        from docling.document_converter import PdfFormatOption
        from docling.datamodel.pipeline_options import PdfPipelineOptions

        # Configure Docling - skip table structure in fast mode for ~3-5x speedup
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = False  # Disable OCR for faster processing
        pipeline_options.do_table_structure = not self.fast_mode  # Skip in fast mode

        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )

        result = converter.convert(str(file_path))
        doc = result.document

        # Extract segments with page numbers
        segments = []
        current_section = None

        for item, level in doc.iterate_items():
            # Track section headers
            if type(item).__name__ == 'SectionHeaderItem':
                current_section = getattr(item, 'text', None)

            # Get text content
            text = getattr(item, 'text', None)
            if not text:
                # For tables, try to get markdown representation
                if type(item).__name__ == 'TableItem':
                    try:
                        text = item.export_to_markdown()
                    except:
                        continue
                else:
                    continue

            # Get page number from provenance
            page_no = None
            if hasattr(item, 'prov') and item.prov:
                prov = item.prov[0]  # Take first provenance item
                if hasattr(prov, 'page_no'):
                    page_no = prov.page_no

            segments.append(TextSegment(
                text=text.strip(),
                page_number=page_no,
                section=current_section
            ))

        # Export as markdown for full content (preserves table structure)
        full_content = doc.export_to_markdown()

        return full_content, segments

    def _load_docx(self, file_path: Path) -> str:
        """Extract text from Word document."""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        return "\n\n".join(paragraphs)

    def _load_text(self, file_path: Path) -> str:
        """Load plain text or markdown file."""
        return file_path.read_text(encoding="utf-8")

    def load_url(self, url: str, use_crawl4ai: bool = True) -> Document:
        """Load document from a URL. Supports PDFs and webpages.

        Args:
            url: The URL to load
            use_crawl4ai: If True, use Crawl4AI for webpages (handles JS, anti-bot).
                          Falls back to BeautifulSoup if Crawl4AI fails.
        """
        parsed = urlparse(url)
        if not parsed.scheme in ("http", "https"):
            raise ValueError(f"Invalid URL scheme: {parsed.scheme}")

        # First, do a HEAD request to check content type without downloading
        # This helps us route PDFs/docs correctly without loading them in Playwright
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,application/pdf,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Connection": "keep-alive",
        }

        path_lower = parsed.path.lower()

        # Check if URL clearly points to a non-HTML document
        is_likely_document = (
            path_lower.endswith(".pdf") or
            path_lower.endswith(".docx") or
            path_lower.endswith((".md", ".markdown")) or
            path_lower.endswith(".txt")
        )

        # For documents, use the traditional requests-based approach
        # with Playwright fallback for anti-bot protected sites
        if is_likely_document:
            content = None
            content_type = ""

            # Try requests first
            try:
                response = requests.get(url, timeout=60, headers=headers)
                response.raise_for_status()
                content = response.content
                content_type = response.headers.get("Content-Type", "").lower()
            except requests.exceptions.HTTPError as e:
                # If we get 403/401, try Playwright as fallback
                if e.response.status_code in (401, 403, 406, 429):
                    print(f"[URL Loader] Got {e.response.status_code}, trying Playwright for document download")
                    content = self._download_with_playwright(url)
                else:
                    raise

            if content is None:
                raise Exception(f"Failed to download document from {url}")

            content_start = content[:8] if len(content) >= 8 else content
            is_pdf_content = content_start.startswith(b'%PDF')

            if is_pdf_content or "application/pdf" in content_type or path_lower.endswith(".pdf"):
                return self._load_pdf_from_bytes(content, url)
            elif path_lower.endswith(".docx"):
                return self._load_docx_from_bytes(content, url)
            elif path_lower.endswith((".md", ".markdown")):
                return Document(
                    content=content.decode('utf-8'),
                    source=url,
                    doc_type="markdown",
                    metadata={"url": url}
                )
            elif path_lower.endswith(".txt"):
                return Document(
                    content=content.decode('utf-8'),
                    source=url,
                    doc_type="text",
                    metadata={"url": url}
                )

        # For webpages, try Crawl4AI first (handles JS, bypasses anti-bot)
        if use_crawl4ai:
            try:
                print(f"[URL Loader] Trying Crawl4AI for {url}")
                return self._load_webpage_with_crawl4ai(url)
            except Exception as e:
                print(f"[URL Loader] Crawl4AI failed, falling back to BeautifulSoup: {e}")

        # Fallback: use requests + BeautifulSoup
        response = requests.get(url, timeout=60, headers=headers)
        response.raise_for_status()

        content_type = response.headers.get("Content-Type", "").lower()
        content_start = response.content[:8] if len(response.content) >= 8 else response.content
        is_pdf_content = content_start.startswith(b'%PDF')

        # Check if we got a PDF even though URL didn't indicate it
        if is_pdf_content or "application/pdf" in content_type:
            return self._load_pdf_from_bytes(response.content, url)

        # Parse as HTML
        return self._load_webpage(response.text, url)

    def _load_pdf_from_bytes(self, content: bytes, source: str) -> Document:
        """Extract text from PDF bytes using Docling for better table/image handling."""
        try:
            return self._load_pdf_from_bytes_with_docling(content, source)
        except Exception as e:
            print(f"Docling failed for URL PDF, falling back to pypdf: {e}")
            return self._load_pdf_from_bytes_basic(content, source)

    def _load_pdf_from_bytes_basic(self, content: bytes, source: str) -> Document:
        """Basic PDF extraction from bytes using pypdf (fallback)."""
        from pypdf import PdfReader
        import io

        reader = PdfReader(io.BytesIO(content))
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return Document(
            content="\n\n".join(text_parts),
            source=source,
            doc_type="pdf",
            metadata={"url": source}
        )

    def _load_pdf_from_bytes_with_docling(self, content: bytes, source: str) -> Document:
        """Extract text from PDF bytes using Docling - handles tables and images."""
        from docling.document_converter import DocumentConverter
        from docling.datamodel.base_models import InputFormat
        from docling.document_converter import PdfFormatOption
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.datamodel.document import DocumentStream
        import tempfile
        import os

        # Configure Docling - skip table structure in fast mode for ~3-5x speedup
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = False  # Disable OCR for faster processing
        pipeline_options.do_table_structure = not self.fast_mode  # Skip in fast mode

        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )

        # Save to temp file since Docling needs a file path
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            result = converter.convert(tmp_path)
            # Export as markdown - this preserves table structure
            markdown_content = result.document.export_to_markdown()
        finally:
            os.unlink(tmp_path)

        return Document(
            content=markdown_content,
            source=source,
            doc_type="pdf",
            metadata={"url": source}
        )

    def _load_docx_from_bytes(self, content: bytes, source: str) -> Document:
        """Extract text from DOCX bytes."""
        from docx import Document as DocxDocument
        import io

        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        return Document(
            content="\n\n".join(paragraphs),
            source=source,
            doc_type="docx",
            metadata={"url": source}
        )

    def _load_webpage(self, html: str, source: str) -> Document:
        """Extract text content from HTML webpage using BeautifulSoup (fallback)."""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()

        # Get text and clean it up
        text = soup.get_text(separator="\n")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        content = "\n".join(lines)

        # Get title if available
        title = soup.title.string if soup.title else None

        return Document(
            content=content,
            source=source,
            doc_type="webpage",
            metadata={"url": source, "title": title}
        )

    def _download_with_playwright(self, url: str) -> bytes:
        """Download a file using curl_cffi with browser TLS impersonation.

        Used as fallback when requests.get() gets 403/401 errors.
        curl_cffi impersonates real browser TLS fingerprints to bypass anti-bot.
        """
        from curl_cffi import requests as curl_requests

        print(f"[Anti-bot] Downloading {url} with Chrome TLS impersonation...")

        # curl_cffi can impersonate real browser TLS fingerprints
        # This bypasses TLS fingerprinting that detects Python requests
        response = curl_requests.get(
            url,
            impersonate="chrome120",  # Impersonate Chrome 120
            timeout=120,
            headers={
                "Accept": "application/pdf,application/octet-stream,*/*;q=0.8",
                "Accept-Language": "en-US,en;q=0.9",
            }
        )

        if response.status_code >= 400:
            raise Exception(f"curl_cffi got {response.status_code}")

        print(f"[Anti-bot] Successfully downloaded {len(response.content)} bytes")
        return response.content

    def _load_webpage_with_crawl4ai(self, url: str) -> Document:
        """Extract text content from webpage using Crawl4AI (handles JS, anti-bot).

        Crawl4AI uses Playwright to render JavaScript and has stealth features
        to bypass common anti-scraping mechanisms.
        """
        import asyncio

        async def crawl():
            from crawl4ai import AsyncWebCrawler
            from crawl4ai.async_configs import BrowserConfig, CrawlerRunConfig

            browser_config = BrowserConfig(
                headless=True,
                verbose=False,
            )
            run_config = CrawlerRunConfig(
                word_count_threshold=10,  # Filter out very short content blocks
                exclude_external_links=True,
                remove_overlay_elements=True,  # Remove popups/modals
                process_iframes=False,  # Skip iframes for speed
            )

            async with AsyncWebCrawler(config=browser_config) as crawler:
                result = await crawler.arun(url=url, config=run_config)

                if not result.success:
                    raise Exception(f"Crawl4AI failed: {result.error_message}")

                return result

        # Run the async crawler
        result = asyncio.run(crawl())

        # Extract title from metadata if available
        title = None
        if hasattr(result, 'metadata') and result.metadata:
            title = result.metadata.get('title')

        # Use markdown content (cleaner for RAG) or fall back to cleaned HTML
        content = result.markdown if result.markdown else result.cleaned_html

        if not content or len(content.strip()) < 50:
            raise Exception("Crawl4AI returned empty or very short content")

        return Document(
            content=content,
            source=url,
            doc_type="webpage",
            metadata={"url": url, "title": title, "crawler": "crawl4ai"}
        )

    def load_git_repository(
        self,
        repo_path: str,
        repo_url: str | None = None,
        commit_hash: str | None = None
    ) -> list[Document]:
        """Load all indexable text files from a cloned git repository.

        Args:
            repo_path: Path to the cloned repository root
            repo_url: Original repository URL (for generating GitHub links)
            commit_hash: Commit hash at time of cloning (for GitHub links)

        Returns:
            List of Document objects for each text file found
        """
        documents = []

        # Convert git clone URL to browse URL for GitHub links
        github_base_url = None
        if repo_url:
            github_base_url = self._git_url_to_browse_url(repo_url, commit_hash)

        # Binary file extensions to skip
        BINARY_EXTENSIONS = {
            '.png', '.jpg', '.jpeg', '.gif', '.ico', '.svg', '.webp', '.bmp', '.tiff',
            '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
            '.zip', '.tar', '.gz', '.rar', '.7z', '.bz2', '.xz',
            '.exe', '.dll', '.so', '.dylib', '.a', '.lib',
            '.mp3', '.mp4', '.wav', '.avi', '.mov', '.mkv', '.flv', '.wmv',
            '.ttf', '.otf', '.woff', '.woff2', '.eot',
            '.pyc', '.pyo', '.class', '.o', '.obj',
            '.db', '.sqlite', '.sqlite3',
            '.lock', '.bin', '.dat', '.pack', '.idx',
            '.jar', '.war', '.ear',
            '.min.js', '.min.css',  # Minified files
        }

        # Directories to skip
        SKIP_DIRS = {
            '.git', 'node_modules', '__pycache__', '.venv', 'venv', 'env',
            'dist', 'build', '.next', '.cache', 'coverage', '.nyc_output',
            '.idea', '.vscode', '.DS_Store', '.pytest_cache', '.mypy_cache',
            'vendor', 'target', 'out', 'bin', 'obj',
            '.eggs', '*.egg-info', '.tox', '.nox',
        }

        for root, dirs, files in os.walk(repo_path):
            # Skip ignored directories (modify in-place to prevent descent)
            dirs[:] = [d for d in dirs if d not in SKIP_DIRS and not d.startswith('.')]

            for filename in files:
                file_path = os.path.join(root, filename)
                ext = os.path.splitext(filename)[1].lower()

                # Skip binary files
                if ext in BINARY_EXTENSIONS:
                    continue

                # Skip hidden files
                if filename.startswith('.'):
                    continue

                # Skip very large files (> 1MB)
                try:
                    if os.path.getsize(file_path) > 1_000_000:
                        continue
                except OSError:
                    continue

                try:
                    # Try to read as text
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()

                    # Skip empty files
                    if not content.strip():
                        continue

                    # Get relative path from repo root
                    rel_path = os.path.relpath(file_path, repo_path)

                    metadata = {
                        'filename': filename,
                        'file_path': rel_path,
                        'repository': True,
                    }
                    # Add GitHub URL info if available
                    if github_base_url:
                        metadata['github_base_url'] = github_base_url

                    documents.append(Document(
                        content=content,
                        source=rel_path,  # Use relative path as source
                        doc_type=self._detect_doc_type_from_ext(ext, filename),
                        metadata=metadata
                    ))
                except (UnicodeDecodeError, IOError):
                    # Skip files that can't be read as text
                    continue

        return documents

    def _detect_doc_type_from_ext(self, ext: str, filename: str) -> str:
        """Detect document type based on file extension."""
        ext_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.tsx': 'typescript',
            '.jsx': 'javascript',
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.rst': 'restructuredtext',
            '.txt': 'text',
            '.json': 'json',
            '.yaml': 'yaml',
            '.yml': 'yaml',
            '.toml': 'toml',
            '.xml': 'xml',
            '.html': 'html',
            '.css': 'css',
            '.scss': 'scss',
            '.sass': 'sass',
            '.less': 'less',
            '.go': 'go',
            '.rs': 'rust',
            '.java': 'java',
            '.kt': 'kotlin',
            '.scala': 'scala',
            '.rb': 'ruby',
            '.php': 'php',
            '.c': 'c',
            '.cpp': 'cpp',
            '.h': 'c',
            '.hpp': 'cpp',
            '.cs': 'csharp',
            '.swift': 'swift',
            '.m': 'objective-c',
            '.sql': 'sql',
            '.sh': 'shell',
            '.bash': 'shell',
            '.zsh': 'shell',
            '.ps1': 'powershell',
            '.dockerfile': 'dockerfile',
            '.r': 'r',
            '.jl': 'julia',
            '.lua': 'lua',
            '.pl': 'perl',
            '.ex': 'elixir',
            '.exs': 'elixir',
            '.erl': 'erlang',
            '.hs': 'haskell',
            '.clj': 'clojure',
            '.vue': 'vue',
            '.svelte': 'svelte',
        }

        # Special filenames
        name_lower = filename.lower()
        if name_lower == 'dockerfile':
            return 'dockerfile'
        if name_lower == 'makefile':
            return 'makefile'
        if name_lower in ('readme', 'readme.md', 'readme.txt'):
            return 'markdown' if ext in ('.md', '.markdown') else 'text'

        return ext_map.get(ext, 'text')

    def _git_url_to_browse_url(self, git_url: str, commit_hash: str | None = None) -> str | None:
        """Convert a git clone URL to a browse URL for viewing files.

        Examples:
            https://github.com/user/repo.git -> https://github.com/user/repo/blob/{commit}/
            https://github.com/user/repo     -> https://github.com/user/repo/blob/{commit}/
            git@github.com:user/repo.git     -> https://github.com/user/repo/blob/{commit}/

        Returns:
            Base URL for file browsing (file path and line numbers added later), or None if unsupported
        """
        import re

        # Strip trailing .git if present
        url = git_url.rstrip('/').removesuffix('.git')

        # Handle SSH URLs (git@github.com:user/repo)
        ssh_match = re.match(r'^git@([^:]+):(.+)$', url)
        if ssh_match:
            host = ssh_match.group(1)
            path = ssh_match.group(2)
            url = f'https://{host}/{path}'

        # Parse the URL to get host and path
        parsed = urlparse(url)
        host = parsed.netloc.lower()
        path = parsed.path.rstrip('/')

        # Determine the ref to use (commit hash or main)
        ref = commit_hash or 'HEAD'

        # GitHub, GitLab, Bitbucket all use similar URL patterns
        if 'github.com' in host:
            return f'{url}/blob/{ref}'
        elif 'gitlab.com' in host:
            return f'{url}/-/blob/{ref}'
        elif 'bitbucket.org' in host:
            return f'{url}/src/{ref}'

        # Unsupported host
        return None
