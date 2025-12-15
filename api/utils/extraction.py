"""Resource extraction utilities for Stage 2 processing.

These functions extract type-specific metadata from uploaded files
without performing semantic enrichment (RAG indexing, LLM summaries).
"""

import json
import mimetypes
from pathlib import Path


def detect_mime_type(content: bytes, filename: str) -> str:
    """Detect MIME type from content and filename.

    Args:
        content: File content bytes (first 2KB is usually enough)
        filename: Original filename

    Returns:
        MIME type string (e.g., "application/pdf")
    """
    # Try python-magic if available
    try:
        import magic
        mime = magic.from_buffer(content[:2048], mime=True)
        if mime and mime != "application/octet-stream":
            return mime
    except ImportError:
        pass

    # Fallback to extension-based detection
    mime_type, _ = mimetypes.guess_type(filename)
    if mime_type:
        return mime_type

    # Manual mapping for common types
    ext = Path(filename).suffix.lower()
    ext_map = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".csv": "text/csv",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".json": "application/json",
        ".parquet": "application/vnd.apache.parquet",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".md": "text/markdown",
        ".txt": "text/plain",
        ".py": "text/x-python",
        ".js": "text/javascript",
        ".ts": "text/typescript",
    }
    return ext_map.get(ext, "application/octet-stream")


def extract_document_metadata(file_path: str) -> dict:
    """Extract metadata from RAG-able documents (PDF, DOCX, MD, TXT).

    Args:
        file_path: Path to the document file

    Returns:
        Dict with extracted metadata (page_count, word_count, etc.)
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf_metadata(file_path)
    elif ext == ".docx":
        return _extract_docx_metadata(file_path)
    elif ext in (".md", ".markdown", ".txt"):
        return _extract_text_metadata(file_path)
    else:
        # Generic text extraction attempt
        return _extract_text_metadata(file_path)


def _extract_pdf_metadata(file_path: str) -> dict:
    """Extract metadata from PDF files."""
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        page_count = len(reader.pages)

        # Extract text and count words (sample first 10 pages for speed)
        word_count = 0
        has_images = False
        sample_pages = min(10, page_count)

        for i in range(sample_pages):
            page = reader.pages[i]
            text = page.extract_text() or ""
            word_count += len(text.split())
            if page.images:
                has_images = True

        # Extrapolate word count if we only sampled
        if sample_pages < page_count:
            word_count = int(word_count * (page_count / sample_pages))

        return {
            "page_count": page_count,
            "word_count": word_count,
            "has_images": has_images,
            "has_tables": False,  # Would need more sophisticated detection
            "pdf_version": reader.pdf_header if hasattr(reader, 'pdf_header') else None
        }
    except Exception as e:
        return {"extraction_error": str(e)}


def _extract_docx_metadata(file_path: str) -> dict:
    """Extract metadata from DOCX files."""
    try:
        import docx
        doc = docx.Document(file_path)

        word_count = 0
        paragraph_count = len(doc.paragraphs)

        for para in doc.paragraphs:
            word_count += len(para.text.split())

        return {
            "word_count": word_count,
            "paragraph_count": paragraph_count,
            "has_tables": len(doc.tables) > 0,
            "table_count": len(doc.tables),
            "page_count": None  # DOCX doesn't have fixed pages
        }
    except Exception as e:
        return {"extraction_error": str(e)}


def _extract_text_metadata(file_path: str) -> dict:
    """Extract metadata from plain text files (MD, TXT, etc.)."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        lines = content.split("\n")
        words = content.split()

        return {
            "word_count": len(words),
            "line_count": len(lines),
            "char_count": len(content),
            "page_count": None
        }
    except Exception as e:
        return {"extraction_error": str(e)}


def extract_data_metadata(file_path: str) -> dict:
    """Extract metadata from data files (CSV, Excel, JSON, Parquet).

    Args:
        file_path: Path to the data file

    Returns:
        Dict with schema info (columns, row_count, dtypes, sample values)
    """
    import pandas as pd
    ext = Path(file_path).suffix.lower()

    try:
        # Read with row limit for large files
        max_rows = 10000

        if ext == ".csv":
            df = pd.read_csv(file_path, nrows=max_rows)
            total_rows = _count_csv_rows(file_path)
        elif ext == ".tsv":
            df = pd.read_csv(file_path, sep="\t", nrows=max_rows)
            total_rows = _count_csv_rows(file_path)
        elif ext in (".xlsx", ".xls"):
            xl = pd.ExcelFile(file_path)
            df = pd.read_excel(xl, nrows=max_rows)
            sheet_names = xl.sheet_names
            # For Excel, use sample row count
            total_rows = len(df)  # Can't easily count without loading
        elif ext == ".json":
            df = pd.read_json(file_path)
            total_rows = len(df)
            df = df.head(max_rows)
        elif ext == ".parquet":
            df = pd.read_parquet(file_path)
            total_rows = len(df)
            df = df.head(max_rows)
        else:
            # Try CSV as fallback
            df = pd.read_csv(file_path, nrows=max_rows)
            total_rows = len(df)

        # Build column schema
        columns = []
        for col in df.columns:
            col_info = {
                "name": str(col),
                "dtype": str(df[col].dtype),
                "null_count": int(df[col].isnull().sum()),
                "null_pct": round(df[col].isnull().sum() / len(df) * 100, 1) if len(df) > 0 else 0
            }
            # Add sample values (non-null, unique)
            non_null = df[col].dropna()
            if len(non_null) > 0:
                samples = non_null.head(5).tolist()
                # Convert to JSON-serializable types
                col_info["sample_values"] = [
                    str(v) if not isinstance(v, (int, float, bool, str, type(None))) else v
                    for v in samples
                ]
            columns.append(col_info)

        result = {
            "row_count": total_rows,
            "column_count": len(df.columns),
            "columns": columns,
            "sample_rows": _df_to_sample_rows(df, n=5)
        }

        # Add sheet names for Excel files
        if ext in (".xlsx", ".xls"):
            result["sheet_names"] = sheet_names

        return result

    except Exception as e:
        return {"extraction_error": str(e)}


def _count_csv_rows(file_path: str) -> int:
    """Count rows in CSV without loading entire file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return sum(1 for _ in f) - 1  # Subtract header
    except Exception:
        return -1


def _df_to_sample_rows(df, n: int = 5) -> list[dict]:
    """Convert first n rows of DataFrame to list of dicts."""
    try:
        sample = df.head(n)
        rows = []
        for _, row in sample.iterrows():
            row_dict = {}
            for col in df.columns:
                val = row[col]
                # Convert to JSON-serializable type
                if pd.isna(val):
                    row_dict[str(col)] = None
                elif isinstance(val, (int, float, bool, str)):
                    row_dict[str(col)] = val
                else:
                    row_dict[str(col)] = str(val)
            rows.append(row_dict)
        return rows
    except Exception:
        return []


def extract_image_metadata(file_path: str) -> dict:
    """Extract metadata from image files (PNG, JPG, GIF, etc.).

    Args:
        file_path: Path to the image file

    Returns:
        Dict with dimensions, format, and mode
    """
    try:
        from PIL import Image

        with Image.open(file_path) as img:
            return {
                "width": img.width,
                "height": img.height,
                "format": img.format,
                "mode": img.mode,  # RGB, RGBA, L, etc.
                "dimensions": f"{img.width}x{img.height}"
            }
    except Exception as e:
        return {"extraction_error": str(e)}


def is_extraction_successful(metadata: dict) -> bool:
    """Check if extraction metadata indicates success."""
    return "extraction_error" not in metadata
